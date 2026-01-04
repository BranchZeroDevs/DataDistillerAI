"""Document loading and parsing utilities."""

from typing import List, Dict, Any
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class Document:
    """Represents a loaded document."""
    
    def __init__(self, content: str, metadata: Dict[str, Any]):
        self.content = content
        self.metadata = metadata
    
    def __repr__(self) -> str:
        return f"Document(source={self.metadata.get('source', 'unknown')}, length={len(self.content)})"


class DocumentLoader:
    """Base class for document loading from various formats."""
    
    SUPPORTED_FORMATS = {'.txt', '.pdf', '.docx', '.html', '.md'}
    
    def __init__(self):
        """Initialize the document loader."""
        self.loaders = self._init_loaders()
    
    def _init_loaders(self) -> Dict[str, callable]:
        """Initialize format-specific loaders."""
        return {
            '.txt': self._load_txt,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.html': self._load_html,
            '.md': self._load_txt,  # Markdown treated as text
        }
    
    def load(self, file_path: str) -> Document:
        """Load a single document."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {suffix}. Supported: {self.SUPPORTED_FORMATS}")
        
        loader = self.loaders.get(suffix, self._load_txt)
        content = loader(file_path)
        
        metadata = {
            'source': str(path),
            'filename': path.name,
            'format': suffix,
        }
        
        logger.info(f"Loaded document: {path.name} ({len(content)} chars)")
        return Document(content, metadata)
    
    def load_directory(self, dir_path: str) -> List[Document]:
        """Load all documents from a directory."""
        dir_path = Path(dir_path)
        documents = []
        
        for file_path in dir_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    doc = self.load(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")
        
        logger.info(f"Loaded {len(documents)} documents from {dir_path}")
        return documents
    
    def _load_txt(self, file_path: str) -> str:
        """Load plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _load_pdf(self, file_path: str) -> str:
        """Load PDF file."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 not installed. Use: pip install PyPDF2")
        
        text = []
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text())
        
        return '\n'.join(text)
    
    def _load_docx(self, file_path: str) -> str:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed. Use: pip install python-docx")
        
        doc = DocxDocument(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        return '\n'.join(text)
    
    def _load_html(self, file_path: str) -> str:
        """Load HTML file."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 not installed. Use: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()
            text = soup.get_text()
        
        return '\n'.join(line.strip() for line in text.split('\n') if line.strip())
